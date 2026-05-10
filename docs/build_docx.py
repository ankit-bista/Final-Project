#!/usr/bin/env python3
"""Build SecureVault_Documentation.docx from Markdown source."""

from pathlib import Path
import re

SRC = Path(__file__).resolve().parent / "SecureVault_Documentation.md"
OUT = Path(__file__).resolve().parent / "SecureVault_Documentation.docx"


def ensure_docx():
    try:
        import docx  # noqa: F401
    except Exception:
        raise SystemExit(
            "python-docx is required. Install with: python3 -m pip install python-docx"
        )


def write_docx():
    from docx import Document

    text = SRC.read_text(encoding="utf-8")
    doc = Document()

    lines = text.splitlines()
    in_code_block = False

    for raw in lines:
        line = raw.rstrip("\n")

        if line.startswith("```"):
            in_code_block = not in_code_block
            if in_code_block:
                doc.add_paragraph("[Diagram block]", style="Intense Quote")
            continue

        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = "No Spacing"
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            title = m.group(2).strip()
            if level == 1:
                doc.add_heading(title, level=0)
            else:
                doc.add_heading(title, level=level)
            continue

        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue

        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            continue

        doc.add_paragraph(line)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    ensure_docx()
    write_docx()
